import io
import logging
from fastapi import HTTPException

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parses PDF, TXT, and DOCX files into plain text."""

    def parse_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            if doc.page_count == 0:
                raise HTTPException(status_code=400, detail="PDF file is empty (no pages).")
            text_parts = []
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text.strip()}")
            doc.close()
            full_text = "\n\n".join(text_parts)
            if not full_text.strip():
                raise HTTPException(status_code=400, detail="PDF contains no extractable text (may be scanned/image-only).")
            return full_text
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {str(e)}")

    def parse_txt(self, file_bytes: bytes) -> str:
        """Decode plain text file."""
        try:
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    text = file_bytes.decode(encoding)
                    if text.strip():
                        return text
                except UnicodeDecodeError:
                    continue
            raise HTTPException(status_code=400, detail="Could not decode text file with any supported encoding.")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"TXT parsing error: {e}")
            raise HTTPException(status_code=400, detail=f"Failed to parse text file: {str(e)}")

    def parse_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            full_text = "\n\n".join(paragraphs)
            if not full_text.strip():
                raise HTTPException(status_code=400, detail="DOCX file contains no extractable text.")
            return full_text
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {str(e)}")

    def parse(self, filename: str, file_bytes: bytes) -> str:
        """Dispatch parsing based on file extension."""
        if not file_bytes:
            raise HTTPException(status_code=400, detail="File is empty.")
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext == "pdf":
            return self.parse_pdf(file_bytes)
        elif ext == "txt":
            return self.parse_txt(file_bytes)
        elif ext == "docx":
            return self.parse_docx(file_bytes)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type '.{ext}'. Allowed: pdf, txt, docx"
            )


document_parser = DocumentParser()
